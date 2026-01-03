from fastapi import FastAPI, HTTPException, Depends, status
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, EmailStr
from datetime import datetime, timedelta
from typing import Optional, List
from jose import jwt, JWTError
from passlib.context import CryptContext
from sqlalchemy import Column, Integer, String, Float, DateTime, ForeignKey, Boolean
from sqlalchemy.ext.asyncio import create_async_engine, AsyncSession, async_sessionmaker
from sqlalchemy.orm import declarative_base, selectinload
from sqlalchemy.future import select
import pandas as pd
from io import BytesIO
from fastapi.responses import StreamingResponse
from config import settings

# Add these imports at the top of main.py
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# Database setup - PostgreSQL with asyncpg
DATABASE_URL = settings.DATABASE_URL.replace("postgresql://", "postgresql+asyncpg://")

engine = create_async_engine(DATABASE_URL, echo=True, pool_pre_ping=True)
AsyncSessionLocal = async_sessionmaker(
    engine, class_=AsyncSession, expire_on_commit=False
)
Base = declarative_base()

# Security
SECRET_KEY = settings.SECRET_KEY
ALGORITHM = settings.ALGORITHM
ACCESS_TOKEN_EXPIRE_MINUTES = settings.ACCESS_TOKEN_EXPIRE_MINUTES
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
security = HTTPBearer()

app = FastAPI(title="Accounting API", version="1.0.0")

# CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Database Models
class User(Base):
    __tablename__ = "users"
    id = Column(Integer, primary_key=True, index=True)
    username = Column(String, unique=True, index=True)
    email = Column(String, unique=True, index=True)
    hashed_password = Column(String)
    company_name = Column(String)
    is_active = Column(Boolean, default=True)
    created_at = Column(DateTime, default=datetime.utcnow)

class Party(Base):
    __tablename__ = "parties"
    id = Column(Integer, primary_key=True, index=True)
    user_id = Column(Integer, ForeignKey("users.id"))
    name = Column(String, index=True)
    party_type = Column(String)  # customer or supplier
    phone = Column(String)
    email = Column(String)
    address = Column(String)
    balance = Column(Float, default=0.0)
    created_at = Column(DateTime, default=datetime.utcnow)

class SaleEntry(Base):
    __tablename__ = "sales"
    id = Column(Integer, primary_key=True, index=True)
    user_id = Column(Integer, ForeignKey("users.id"))
    party_id = Column(Integer, ForeignKey("parties.id"))
    invoice_no = Column(String, unique=True, index=True)
    date = Column(DateTime)
    item_name = Column(String)
    quantity = Column(Float)
    rate = Column(Float)
    amount = Column(Float)
    tax = Column(Float, default=0.0)
    total_amount = Column(Float)
    payment_status = Column(String, default="pending")  # paid, pending, partial
    notes = Column(String)
    created_at = Column(DateTime, default=datetime.utcnow)

class PurchaseEntry(Base):
    __tablename__ = "purchases"
    id = Column(Integer, primary_key=True, index=True)
    user_id = Column(Integer, ForeignKey("users.id"))
    party_id = Column(Integer, ForeignKey("parties.id"))
    bill_no = Column(String, index=True)
    date = Column(DateTime)
    item_name = Column(String)
    quantity = Column(Float)
    rate = Column(Float)
    amount = Column(Float)
    tax = Column(Float, default=0.0)
    total_amount = Column(Float)
    payment_status = Column(String, default="pending")
    notes = Column(String)
    created_at = Column(DateTime, default=datetime.utcnow)

# Create tables
@app.on_event("startup")
async def startup():
    async with engine.begin() as conn:
        await conn.run_sync(Base.metadata.create_all)

# Pydantic Models
class UserCreate(BaseModel):
    username: str
    email: EmailStr
    password: str
    company_name: str

class UserLogin(BaseModel):
    username: str
    password: str

class PartyCreate(BaseModel):
    name: str
    party_type: str
    phone: Optional[str] = None
    email: Optional[str] = None
    address: Optional[str] = None

class SaleCreate(BaseModel):
    party_id: int
    invoice_no: str
    date: datetime
    item_name: str
    quantity: float
    rate: float
    tax: float = 0.0
    payment_status: str = "pending"
    notes: Optional[str] = None

class PurchaseCreate(BaseModel):
    party_id: int
    bill_no: str
    date: datetime
    item_name: str
    quantity: float
    rate: float
    tax: float = 0.0
    payment_status: str = "pending"
    notes: Optional[str] = None

# Dependency
async def get_db():
    async with AsyncSessionLocal() as session:
        yield session

# Authentication Functions
def create_access_token(data: dict):
    to_encode = data.copy()
    expire = datetime.utcnow() + timedelta(days=30)
    to_encode.update({"exp": expire})
    return jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)

def verify_token(credentials: HTTPAuthorizationCredentials = Depends(security)):
    try:
        token = credentials.credentials
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        return payload
    except:
        raise HTTPException(status_code=401, detail="Invalid token")

async def get_current_user(
    token_data: dict = Depends(verify_token), 
    db: AsyncSession = Depends(get_db)
):
    result = await db.execute(select(User).filter(User.id == token_data["user_id"]))
    user = result.scalar_one_or_none()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    return user

# API Endpoints
@app.post("/register")
async def register(user: UserCreate, db: AsyncSession = Depends(get_db)):
    result = await db.execute(select(User).filter(User.username == user.username))
    if result.scalar_one_or_none():
        raise HTTPException(status_code=400, detail="Username already exists")
    
    hashed_pw = pwd_context.hash(user.password)
    new_user = User(
        username=user.username,
        email=user.email,
        hashed_password=hashed_pw,
        company_name=user.company_name
    )
    db.add(new_user)
    await db.commit()
    await db.refresh(new_user)
    
    token = create_access_token({"user_id": new_user.id})
    return {"token": token, "user_id": new_user.id, "username": new_user.username}

@app.post("/login")
async def login(user: UserLogin, db: AsyncSession = Depends(get_db)):
    result = await db.execute(select(User).filter(User.username == user.username))
    db_user = result.scalar_one_or_none()
    
    if not db_user or not pwd_context.verify(user.password, db_user.hashed_password):
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    token = create_access_token({"user_id": db_user.id})
    return {"token": token, "user_id": db_user.id, "username": db_user.username}

@app.get("/dashboard")
async def get_dashboard(
    current_user: User = Depends(get_current_user), 
    db: AsyncSession = Depends(get_db)
):
    sales_result = await db.execute(
        select(SaleEntry).filter(SaleEntry.user_id == current_user.id)
    )
    sales = sales_result.scalars().all()
    
    purchases_result = await db.execute(
        select(PurchaseEntry).filter(PurchaseEntry.user_id == current_user.id)
    )
    purchases = purchases_result.scalars().all()
    
    total_sales = sum(s.total_amount for s in sales)
    total_purchases = sum(p.total_amount for p in purchases)
    profit = total_sales - total_purchases
    
    return {
        "total_sales": total_sales,
        "total_purchases": total_purchases,
        "profit": profit,
        "sales_count": len(sales),
        "purchases_count": len(purchases)
    }

@app.post("/parties")
async def create_party(
    party: PartyCreate, 
    current_user: User = Depends(get_current_user), 
    db: AsyncSession = Depends(get_db)
):
    new_party = Party(**party.dict(), user_id=current_user.id)
    db.add(new_party)
    await db.commit()
    await db.refresh(new_party)
    return new_party

@app.get("/parties")
async def get_parties(
    party_type: Optional[str] = None, 
    current_user: User = Depends(get_current_user), 
    db: AsyncSession = Depends(get_db)
):
    query = select(Party).filter(Party.user_id == current_user.id)
    if party_type:
        query = query.filter(Party.party_type == party_type)
    
    result = await db.execute(query)
    return result.scalars().all()

@app.post("/sales")
async def create_sale(
    sale: SaleCreate, 
    current_user: User = Depends(get_current_user), 
    db: AsyncSession = Depends(get_db)
):
    amount = sale.quantity * sale.rate
    total_amount = amount + sale.tax
    
    new_sale = SaleEntry(
        **sale.dict(),
        user_id=current_user.id,
        amount=amount,
        total_amount=total_amount
    )
    db.add(new_sale)
    await db.commit()
    await db.refresh(new_sale)
    return new_sale

@app.get("/sales")
async def get_sales(
    current_user: User = Depends(get_current_user), 
    db: AsyncSession = Depends(get_db)
):
    result = await db.execute(
        select(SaleEntry).filter(SaleEntry.user_id == current_user.id)
    )
    return result.scalars().all()

@app.post("/purchases")
async def create_purchase(
    purchase: PurchaseCreate, 
    current_user: User = Depends(get_current_user), 
    db: AsyncSession = Depends(get_db)
):
    amount = purchase.quantity * purchase.rate
    total_amount = amount + purchase.tax
    
    new_purchase = PurchaseEntry(
        **purchase.dict(),
        user_id=current_user.id,
        amount=amount,
        total_amount=total_amount
    )
    db.add(new_purchase)
    await db.commit()
    await db.refresh(new_purchase)
    return new_purchase

@app.get("/purchases")
async def get_purchases(
    current_user: User = Depends(get_current_user), 
    db: AsyncSession = Depends(get_db)
):
    result = await db.execute(
        select(PurchaseEntry).filter(PurchaseEntry.user_id == current_user.id)
    )
    return result.scalars().all()

@app.get("/export/sales/excel")
async def export_sales_excel(
    current_user: User = Depends(get_current_user), 
    db: AsyncSession = Depends(get_db)
):
    result = await db.execute(
        select(SaleEntry).filter(SaleEntry.user_id == current_user.id)
    )
    sales = result.scalars().all()
    
    data = [{
        "Invoice No": s.invoice_no,
        "Date": s.date.strftime("%Y-%m-%d"),
        "Item": s.item_name,
        "Quantity": s.quantity,
        "Rate": s.rate,
        "Amount": s.amount,
        "Tax": s.tax,
        "Total": s.total_amount,
        "Status": s.payment_status
    } for s in sales]
    
    df = pd.DataFrame(data)
    output = BytesIO()
    df.to_excel(output, index=False, engine='openpyxl')
    output.seek(0)
    
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=sales_report.xlsx"}
    )

@app.get("/export/purchases/excel")
async def export_purchases_excel(
    current_user: User = Depends(get_current_user), 
    db: AsyncSession = Depends(get_db)
):
    result = await db.execute(
        select(PurchaseEntry).filter(PurchaseEntry.user_id == current_user.id)
    )
    purchases = result.scalars().all()
    
    data = [{
        "Bill No": p.bill_no,
        "Date": p.date.strftime("%Y-%m-%d"),
        "Item": p.item_name,
        "Quantity": p.quantity,
        "Rate": p.rate,
        "Amount": p.amount,
        "Tax": p.tax,
        "Total": p.total_amount,
        "Status": p.payment_status
    } for p in purchases]
    
    df = pd.DataFrame(data)
    output = BytesIO()
    df.to_excel(output, index=False, engine='openpyxl')
    output.seek(0)
    
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=purchases_report.xlsx"}
    )

#pdf and docx export endpoints can be added similarly
@app.get("/export/sales/pdf")
def export_sales_pdf(current_user: User = Depends(get_current_user), db: AsyncSession = Depends(get_db)):
    sales = db.query(SaleEntry).filter(SaleEntry.user_id == current_user.id).all()
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    elements = []
    
    # Styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#2196F3'),
        spaceAfter=30,
        alignment=1  # Center
    )
    
    # Title
    title = Paragraph(f"Sales Report - {current_user.company_name}", title_style)
    elements.append(title)
    elements.append(Spacer(1, 0.3*inch))
    
    # Summary
    total_sales = sum(s.total_amount for s in sales)
    summary_style = styles['Normal']
    summary = Paragraph(f"<b>Total Sales:</b> ₹{total_sales:,.2f}<br/><b>Total Transactions:</b> {len(sales)}", summary_style)
    elements.append(summary)
    elements.append(Spacer(1, 0.3*inch))
    
    # Table Data
    data = [['Invoice', 'Date', 'Item', 'Qty', 'Rate', 'Tax', 'Total', 'Status']]
    
    for sale in sales:
        data.append([
            sale.invoice_no,
            sale.date.strftime("%d-%m-%Y"),
            sale.item_name[:20],  # Truncate long names
            str(sale.quantity),
            f"₹{sale.rate:.2f}",
            f"₹{sale.tax:.2f}",
            f"₹{sale.total_amount:.2f}",
            sale.payment_status.upper()
        ])
    
    # Create Table
    table = Table(data, repeatRows=1)
    table.setStyle(TableStyle([
        # Header Style
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2196F3')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        
        # Data Style
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 8),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#E3F2FD')]),
        ('TOPPADDING', (0, 1), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 8),
    ]))
    
    elements.append(table)
    
    # Footer
    elements.append(Spacer(1, 0.5*inch))
    footer = Paragraph(
        f"Generated on: {datetime.now().strftime('%d %B %Y, %I:%M %p')}<br/>Powered by Accounting Pro",
        styles['Normal']
    )
    elements.append(footer)
    
    # Build PDF
    doc.build(elements)
    buffer.seek(0)
    
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": "attachment; filename=sales_report.pdf"}
    )


@app.get("/export/purchases/pdf")
def export_purchases_pdf(current_user: User = Depends(get_current_user), db: AsyncSession = Depends(get_db)):
    purchases = db.query(PurchaseEntry).filter(PurchaseEntry.user_id == current_user.id).all()
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    elements = []
    
    # Styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#FF9800'),
        spaceAfter=30,
        alignment=1
    )
    
    # Title
    title = Paragraph(f"Purchase Report - {current_user.company_name}", title_style)
    elements.append(title)
    elements.append(Spacer(1, 0.3*inch))
    
    # Summary
    total_purchases = sum(p.total_amount for p in purchases)
    summary_style = styles['Normal']
    summary = Paragraph(f"<b>Total Purchases:</b> ₹{total_purchases:,.2f}<br/><b>Total Transactions:</b> {len(purchases)}", summary_style)
    elements.append(summary)
    elements.append(Spacer(1, 0.3*inch))
    
    # Table Data
    data = [['Bill No', 'Date', 'Item', 'Qty', 'Rate', 'Tax', 'Total', 'Status']]
    
    for purchase in purchases:
        data.append([
            purchase.bill_no,
            purchase.date.strftime("%d-%m-%Y"),
            purchase.item_name[:20],
            str(purchase.quantity),
            f"₹{purchase.rate:.2f}",
            f"₹{purchase.tax:.2f}",
            f"₹{purchase.total_amount:.2f}",
            purchase.payment_status.upper()
        ])
    
    # Create Table
    table = Table(data, repeatRows=1)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#FF9800')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 8),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#FFF3E0')]),
        ('TOPPADDING', (0, 1), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 8),
    ]))
    
    elements.append(table)
    elements.append(Spacer(1, 0.5*inch))
    footer = Paragraph(
        f"Generated on: {datetime.now().strftime('%d %B %Y, %I:%M %p')}<br/>Powered by Accounting Pro",
        styles['Normal']
    )
    elements.append(footer)
    
    doc.build(elements)
    buffer.seek(0)
    
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": "attachment; filename=purchases_report.pdf"}
    )


@app.get("/export/sales/word")
def export_sales_word(current_user: User = Depends(get_current_user), db: AsyncSession = Depends(get_db)):
    sales = db.query(SaleEntry).filter(SaleEntry.user_id == current_user.id).all()
    
    document = Document()
    
    # Title
    title = document.add_heading(f'Sales Report - {current_user.company_name}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Summary
    document.add_paragraph()
    summary = document.add_paragraph()
    summary.add_run('Total Sales: ').bold = True
    summary.add_run(f'₹{sum(s.total_amount for s in sales):,.2f}')
    summary.add_run('\nTotal Transactions: ').bold = True
    summary.add_run(f'{len(sales)}')
    
    document.add_paragraph()
    
    # Table
    table = document.add_table(rows=1, cols=8)
    table.style = 'Light Grid Accent 1'
    
    # Header
    header_cells = table.rows[0].cells
    headers = ['Invoice', 'Date', 'Item', 'Qty', 'Rate', 'Tax', 'Total', 'Status']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        # Make header bold
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)
    
    # Data rows
    for sale in sales:
        row_cells = table.add_row().cells
        row_cells[0].text = sale.invoice_no
        row_cells[1].text = sale.date.strftime("%d-%m-%Y")
        row_cells[2].text = sale.item_name
        row_cells[3].text = str(sale.quantity)
        row_cells[4].text = f'₹{sale.rate:.2f}'
        row_cells[5].text = f'₹{sale.tax:.2f}'
        row_cells[6].text = f'₹{sale.total_amount:.2f}'
        row_cells[7].text = sale.payment_status.upper()
    
    # Footer
    document.add_paragraph()
    footer = document.add_paragraph()
    footer.add_run(f'Generated on: {datetime.now().strftime("%d %B %Y, %I:%M %p")}').italic = True
    footer.add_run('\nPowered by Accounting Pro').italic = True
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save to buffer
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=sales_report.docx"}
    )


@app.get("/export/purchases/word")
def export_purchases_word(current_user: User = Depends(get_current_user), db: AsyncSession = Depends(get_db)):
    purchases = db.query(PurchaseEntry).filter(PurchaseEntry.user_id == current_user.id).all()
    
    document = Document()
    
    # Title
    title = document.add_heading(f'Purchase Report - {current_user.company_name}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Summary
    document.add_paragraph()
    summary = document.add_paragraph()
    summary.add_run('Total Purchases: ').bold = True
    summary.add_run(f'₹{sum(p.total_amount for p in purchases):,.2f}')
    summary.add_run('\nTotal Transactions: ').bold = True
    summary.add_run(f'{len(purchases)}')
    
    document.add_paragraph()
    
    # Table
    table = document.add_table(rows=1, cols=8)
    table.style = 'Light Grid Accent 1'
    
    # Header
    header_cells = table.rows[0].cells
    headers = ['Bill No', 'Date', 'Item', 'Qty', 'Rate', 'Tax', 'Total', 'Status']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)
    
    # Data rows
    for purchase in purchases:
        row_cells = table.add_row().cells
        row_cells[0].text = purchase.bill_no
        row_cells[1].text = purchase.date.strftime("%d-%m-%Y")
        row_cells[2].text = purchase.item_name
        row_cells[3].text = str(purchase.quantity)
        row_cells[4].text = f'₹{purchase.rate:.2f}'
        row_cells[5].text = f'₹{purchase.tax:.2f}'
        row_cells[6].text = f'₹{purchase.total_amount:.2f}'
        row_cells[7].text = purchase.payment_status.upper()
    
    # Footer
    document.add_paragraph()
    footer = document.add_paragraph()
    footer.add_run(f'Generated on: {datetime.now().strftime("%d %B %Y, %I:%M %p")}').italic = True
    footer.add_run('\nPowered by Accounting Pro').italic = True
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save to buffer
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=purchases_report.docx"}
    )


# Bonus: Combined Profit/Loss Report
@app.get("/export/profit-loss/pdf")
def export_profit_loss_pdf(current_user: User = Depends(get_current_user), db: AsyncSession = Depends(get_db)):
    sales = db.query(SaleEntry).filter(SaleEntry.user_id == current_user.id).all()
    purchases = db.query(PurchaseEntry).filter(PurchaseEntry.user_id == current_user.id).all()
    
    total_sales = sum(s.total_amount for s in sales)
    total_purchases = sum(p.total_amount for p in purchases)
    profit = total_sales - total_purchases
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    elements = []
    styles = getSampleStyleSheet()
    
    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=28,
        textColor=colors.HexColor('#673AB7'),
        spaceAfter=30,
        alignment=1
    )
    title = Paragraph(f"Profit & Loss Statement<br/>{current_user.company_name}", title_style)
    elements.append(title)
    elements.append(Spacer(1, 0.5*inch))
    
    # Summary Table
    summary_data = [
        ['Description', 'Amount'],
        ['Total Sales Revenue', f'₹{total_sales:,.2f}'],
        ['Total Purchase Cost', f'₹{total_purchases:,.2f}'],
        ['', ''],
        ['Net Profit/Loss', f'₹{profit:,.2f}'],
    ]
    
    summary_table = Table(summary_data, colWidths=[3*inch, 2*inch])
    summary_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#673AB7')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('ALIGN', (1, 0), (1, -1), 'RIGHT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('GRID', (0, 0), (-1, -2), 0.5, colors.grey),
        ('LINEABOVE', (0, -1), (-1, -1), 2, colors.HexColor('#673AB7')),
        ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, -1), (-1, -1), 14),
        ('TOPPADDING', (0, -1), (-1, -1), 10),
    ]))
    
    elements.append(summary_table)
    elements.append(Spacer(1, 0.5*inch))
    
    # Footer
    footer = Paragraph(
        f"Period: All Time<br/>Generated on: {datetime.now().strftime('%d %B %Y, %I:%M %p')}<br/>Powered by Accounting Pro",
        styles['Normal']
    )
    elements.append(footer)
    
    doc.build(elements)
    buffer.seek(0)
    
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": "attachment; filename=profit_loss_statement.pdf"}
    )

# if __name__ == "__main__":
#     import uvicorn
#     uvicorn.run(app, host="0.0.0.0", port=8000)